"""Dify DSL 组装器：把架构师 LLM 的 BlueprintSpec 渲染成可导入的 Dify 0.3.0 app dict。
本文件为纯 stdlib，源码会被原样内嵌进生成器的 Dify 代码节点。"""
import json
import re
import uuid

MODEL_PROVIDER = "csg/ai_model_provider/ai_gateway_platform"
FLASH_MODEL_NAME = "DeepSeek-V4-Flash"
PRO_MODEL_NAME = "DeepSeek-V4-Pro"


def clean_json(raw):
    text = str(raw or "").strip()
    text = re.sub(r"^```(?:json)?", "", text).strip()
    text = re.sub(r"```$", "", text).strip()
    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        text = text[start:end + 1]
    try:
        data = json.loads(text)
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def safe_id(value, fallback):
    value = re.sub(r"[^a-zA-Z0-9_]", "_", str(value or "")).strip("_").lower()
    return value or fallback


def as_list(value):
    if value is None:
        return []
    if isinstance(value, list):
        return [v for v in value if str(v).strip()]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def safe_filename(name):
    raw = str(name or "Dify智能体").strip()
    raw = re.sub(r"[\\/:*?\"<>|]+", "", raw)
    raw = re.sub(r"\s+", "", raw)
    if not raw:
        raw = "Dify智能体"
    if not raw.lower().endswith((".yml", ".yaml")):
        raw += ".yml"
    return raw[:80]


def node(node_id, dtype, title, x, y, data, height=90, width=244):
    return {
        "data": {"desc": "", "selected": False, "title": title, "type": dtype, **data},
        "height": height,
        "id": node_id,
        "position": {"x": x, "y": y},
        "positionAbsolute": {"x": x, "y": y},
        "selected": False,
        "sourcePosition": "right",
        "targetPosition": "left",
        "type": "custom",
        "width": width,
    }


def edge(source, handle, target, source_type, target_type):
    return {
        "data": {"isInLoop": False, "sourceType": source_type, "targetType": target_type},
        "id": f"{source}-{handle}-{target}-target",
        "source": source,
        "sourceHandle": handle,
        "target": target,
        "targetHandle": "target",
        "type": "custom",
        "zIndex": 0,
    }


def model(provider, name, temperature=0.2):
    return {
        "completion_params": {"temperature": temperature},
        "mode": "chat",
        "name": name,
        "provider": provider,
    }


def scalar(value, indent=0):
    if value is None:
        return "null"
    if value is True:
        return "true"
    if value is False:
        return "false"
    if isinstance(value, (int, float)):
        return str(value)
    s = str(value)
    if "\n" in s:
        lines = s.splitlines() or [""]
        block_pad = " " * (indent + 2)
        return "|-\n" + "\n".join(block_pad + line for line in lines)
    return json.dumps(s, ensure_ascii=False)


def dump_yaml(obj, indent=0):
    pad = " " * indent
    if isinstance(obj, dict):
        lines = []
        for k, v in obj.items():
            key = str(k)
            if isinstance(v, (dict, list)):
                lines.append(f"{pad}{key}:")
                lines.append(dump_yaml(v, indent + 2))
            else:
                lines.append(f"{pad}{key}: {scalar(v, indent)}")
        return "\n".join(lines)
    if isinstance(obj, list):
        if not obj:
            return pad + "[]"
        lines = []
        for item in obj:
            if isinstance(item, (dict, list)):
                lines.append(f"{pad}-")
                lines.append(dump_yaml(item, indent + 2))
            else:
                lines.append(f"{pad}- {scalar(item, indent)}")
        return "\n".join(lines)
    return pad + scalar(obj, indent)


def render_knowledge(spec, node_id, x, y, provider):
    n = node(node_id, "knowledge-retrieval", "知识检索", x, y, {
        "dataset_ids": spec.get("dataset_ids") or [],
        "multiple_retrieval_config": {
            "reranking_enable": False,
            "reranking_mode": "reranking_model",
            "reranking_model": {"model": "Qwen3-Reranker-4B", "provider": provider},
            "top_k": 5,
        },
        "query_variable_selector": ["sys", "query"],
        "retrieval_mode": "multiple",
    }, height=92)
    return n, [node_id, "result"]


def render_document_extractor(node_id, x, y, start_id=None):
    # 读取 sys.files(数组)：无文件时为空数组 [] 而非 None，避免“unsupported variable type NoneType”。
    n = node(node_id, "document-extractor", "文件文本抽取", x, y, {
        "desc": "抽取上传文档文本(读取 sys.files，无文件时为空，不报错)。图片请用多模态LLM处理。",
        "is_array_file": True,
        "variable_selector": ["sys", "files"],
    }, height=92)
    return n, [node_id, "text"]


def render_http(spec, node_id, x, y):
    base_env = safe_id(spec.get("base_env"), "API_BASE_URL").upper()
    token_env = safe_id(spec.get("token_env"), "API_TOKEN").upper()
    path = spec.get("path") or "/v1/your-endpoint"
    url = "{{#env." + base_env + "#}}" + path
    auth = str(spec.get("auth") or "bearer").lower()
    body = spec.get("body") or None
    headers = "Accept: application/json"
    if body:
        headers += "\nContent-Type: application/json"
    if auth != "none":
        headers += "\nAuthorization: Bearer {{#env." + token_env + "#}}"
    n = node(node_id, "http-request", "HTTP 请求：" + (spec.get("name") or "业务接口"), x, y, {
        "authorization": {"config": None, "type": "no-auth"},
        "body": {"data": [{"type": "text", "value": body}] if body else [{"type": "text", "value": ""}],
                 "type": "json" if body else "none"},
        "headers": headers,
        "method": spec.get("method") or "GET",
        "params": "",
        "retry_config": {"max_retries": 2, "retry_enabled": True, "retry_interval": 500},
        "ssl_verify": False,
        "timeout": {"max_connect_timeout": 20, "max_read_timeout": 60, "max_write_timeout": 60},
        "url": url,
        "variables": [],
    }, height=156)
    return n, [node_id, "body"]


# 把蓝图里的 input "from" 映射到上游节点输出选择器
_FROM_FIELD = {"prev_llm": "text", "prev_http": "body", "prev_doc": "text", "prev_code": "result"}


def _harden_none(code_text):
    """把 LLM 常写的 str(x).strip() 自动改成 str(x or '').strip()。
    无文件/空返回时 Dify 传 None,str(None)='None' 非空会误判,这里统一兜底。"""
    return re.sub(r"str\(\s*([A-Za-z_]\w*)\s*\)\s*\.strip\(\)", r"str(\1 or '').strip()", str(code_text or ""))


def render_code(spec, node_id, x, y, upstream_refs):
    """upstream_refs: dict 形如 {"prev_llm": node_id, "prev_http": node_id, ...}"""
    lines = as_list(spec.get("python_lines"))
    code_text = "\n".join(lines) if lines else "def main(query: str) -> dict:\n    return {'result': query}"
    code_text = _harden_none(code_text)  # 通用 None 安全网
    desc = "代码节点：" + str(spec.get("purpose") or "数据处理")
    try:
        compile(code_text, "<code-node>", "exec")
    except SyntaxError as exc:
        desc += " | SYNTAX_ERROR: " + str(exc)[:160]
    if "def main" not in code_text:
        desc += " | SYNTAX_ERROR: 缺少 def main(...) 入口函数"
    variables = []
    for item in spec.get("inputs") or []:
        var = safe_id(item.get("variable"), "var")
        src = str(item.get("from") or "query")
        if src == "query":
            variables.append({"value_selector": ["sys", "query"], "variable": var})
        elif src in _FROM_FIELD and upstream_refs.get(src):
            variables.append({"value_selector": [upstream_refs[src], _FROM_FIELD[src]], "variable": var})
    if not variables:
        variables = [{"value_selector": ["sys", "query"], "variable": "query"}]
    # main() 的必填参数必须都在 inputs 里声明,否则 Dify 运行时缺参报错。
    declared = {v["variable"] for v in variables}
    sig = re.search(r"def\s+main\s*\(([^)]*)\)", code_text)
    if sig:
        for part in sig.group(1).split(","):
            part = part.strip()
            if not part or "=" in part or part.startswith("*"):
                continue
            pname = re.split(r"[:\s]", part)[0].strip()
            if pname and pname not in declared:
                desc += " | SYNTAX_ERROR: main 参数 '" + pname + "' 未在 inputs 声明(会运行时缺参)"
    n = node(node_id, "code", "代码处理", x, y, {
        "desc": desc,
        "code": code_text,
        "code_language": "python3",
        "outputs": {"result": {"children": None, "type": "string"}},
        "variables": variables,
    }, height=54)
    return n, [node_id, "result"]


def render_llm(spec, node_id, x, y, provider, model_name, memory, context_ref, upstream=None, start_id=None):
    upstream = upstream or {}
    vision_on = bool(spec.get("vision"))
    # user 消息是【固定接线管道】：永远用运行时的 {{#sys.query#}} + 自动注入的上游上下文。
    # 【不采用】架构师写的 user_template——它常把"生成时的原始需求"硬编码进来,导致生成的智能体收不到终端用户的真实输入。
    user_text = "用户问题：{{#sys.query#}}"
    variables = [{"value_selector": ["sys", "query"], "variable": "query"}]
    appended = []
    # 数据接线由 Python 拥有：把上游真实节点ID注入 user 提示词并声明 variables。
    plumbing = [
        ("prev_knowledge", "result", "knowledge_result", "知识库结果"),
        ("prev_doc", "text", "file_text", "上传文件文本"),
        ("prev_http", "body", "http_body", "HTTP返回"),
        ("prev_code", "result", "code_result", "代码处理结果"),
    ]
    for key, field, varname, label in plumbing:
        ref_id = upstream.get(key)
        if ref_id:
            variables.append({"value_selector": [ref_id, field], "variable": varname})
            appended.append(label + "：{{#" + ref_id + "." + field + "#}}")
    if appended:
        user_text = user_text + "\n" + "\n".join(appended)
    data = {
        "context": {"enabled": bool(context_ref), "variable_selector": context_ref or []},
        "memory": {"query_prompt_template": "{{#sys.query#}}", "role_prefix": {"assistant": "", "user": ""},
                   "window": {"enabled": bool(memory), "size": 8}},
        "model": model(provider, model_name, 0.2),
        "prompt_template": [
            {"id": str(uuid.uuid4()), "role": "system", "text": spec.get("system_prompt") or "你是专业助手。"},
            {"id": str(uuid.uuid4()), "role": "user", "text": user_text},
        ],
        "variables": variables,
        "vision": ({"enabled": True, "configs": {"detail": "high",
                    "variable_selector": [start_id or "node_start", "upload_file"]}}
                   if vision_on else {"enabled": False}),
    }
    n = node(node_id, "llm", "图像识别" if vision_on else "回复生成", x, y, data, height=108)
    return n, [node_id, "text"]


def render_answer(spec, node_id, x, y, resolved_text=None):
    text = resolved_text if resolved_text is not None else (spec.get("template") or "{{#sys.query#}}")
    return node(node_id, "answer", "直接回复", x, y,
                {"answer": text, "variables": []}, height=100)


def file_generation_code(output_format="auto", filename_hint="dify_export"):
    fmt = re.sub(r"[^a-z0-9]", "", str(output_format or "auto").lower()) or "auto"
    hint = re.sub(r"[\\/:*?\"<>|\\s]+", "_", str(filename_hint or "dify_export")).strip("_") or "dify_export"
    template = """import csv
import html
import json
import os
import re
import uuid
from datetime import datetime

DEFAULT_FORMAT = "__DEFAULT_FORMAT__"
FILENAME_HINT = "__FILENAME_HINT__"
EXPORT_DIR = "/tmp/dify_generated_files"

def _safe_name(name):
    name = re.sub(r"[\\\\/:*?\\"<>|\\s]+", "_", str(name or "dify_export")).strip("_")
    return (name or "dify_export")[:80]

def _pick_format(query):
    text = str(query or "").lower()
    if "xlsx" in text or "excel" in text or "表格" in text or "统计表" in text:
        return "xlsx"
    if "csv" in text:
        return "csv"
    if "pdf" in text:
        return "pdf"
    if "html" in text:
        return "html"
    if "docx" in text or "word" in text or "报告" in text or "通知" in text or "总结" in text:
        return "docx"
    return DEFAULT_FORMAT if DEFAULT_FORMAT != "auto" else "docx"

def _clean_text(value):
    text = str(value or "").replace("\\r\\n", "\\n").replace("\\r", "\\n")
    return text.strip() or "未生成正文内容，请补充材料后重新生成。"

def _split_lines(text):
    return [line.strip() for line in str(text or "").split("\\n") if line.strip()]

def _write_docx(path, title, content, structured):
    from docx import Document
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    doc.add_paragraph("")
    for line in _split_lines(content):
        plain = re.sub(r"^#+\\s*", "", line).strip()
        if not plain:
            continue
        if line.startswith("# "):
            doc.add_heading(plain, level=1)
        elif line.startswith("## ") or re.match(r"^[一二三四五六七八九十]+[、.．]", plain):
            doc.add_heading(plain, level=2)
        elif line.startswith(("- ", "* ", "· ")):
            doc.add_paragraph(plain[2:].strip(), style="List Bullet")
        elif re.match(r"^\\d+[).、]\\s*", plain):
            doc.add_paragraph(plain, style="List Number")
        else:
            doc.add_paragraph(plain)
    if structured:
        doc.add_page_break()
        doc.add_heading("结构化处理结果", level=2)
        doc.add_paragraph(str(structured)[:8000])
    doc.save(path)

def _write_xlsx(path, title, content, structured):
    try:
        from openpyxl import Workbook
    except Exception:
        csv_path = re.sub(r"\\.xlsx$", ".csv", path)
        _write_csv(csv_path, title, content, structured)
        return csv_path
    wb = Workbook()
    ws = wb.active
    ws.title = "报告内容"
    ws.append(["标题", title])
    ws.append(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
    ws.append([])
    ws.append(["序号", "内容"])
    for idx, line in enumerate(_split_lines(content), 1):
        ws.append([idx, line])
    ws2 = wb.create_sheet("结构化结果")
    ws2.append(["字段", "值"])
    try:
        data = json.loads(structured) if isinstance(structured, str) and structured.strip().startswith(("{", "[")) else structured
        if isinstance(data, dict):
            for key, value in data.items():
                ws2.append([str(key), json.dumps(value, ensure_ascii=False) if isinstance(value, (dict, list)) else str(value)])
        else:
            ws2.append(["raw", str(structured or "")[:32000]])
    except Exception:
        ws2.append(["raw", str(structured or "")[:32000]])
    for sheet in wb.worksheets:
        for col in sheet.columns:
            width = min(max(len(str(cell.value or "")) for cell in col) + 2, 60)
            sheet.column_dimensions[col[0].column_letter].width = width
    wb.save(path)
    return path

def _write_csv(path, title, content, structured):
    with open(path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["标题", title])
        writer.writerow(["生成时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
        writer.writerow([])
        writer.writerow(["序号", "内容"])
        for idx, line in enumerate(_split_lines(content), 1):
            writer.writerow([idx, line])
        if structured:
            writer.writerow([])
            writer.writerow(["结构化处理结果", str(structured)[:32000]])

def _write_html(path, title, content, structured):
    escaped = "<br/>".join(html.escape(line) for line in _split_lines(content))
    structured_html = html.escape(str(structured or ""))
    body = (
        "<!doctype html>\\n"
        "<html><head><meta charset=\\"utf-8\\"><title>" + html.escape(title) + "</title>"
        "<style>body{font-family:Arial,'Microsoft YaHei',sans-serif;line-height:1.7;margin:40px;}pre{white-space:pre-wrap;background:#f6f8fa;padding:16px;}</style>"
        "</head><body><h1>" + html.escape(title) + "</h1><p>生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M:%S") + "</p><div>" + escaped + "</div><h2>结构化处理结果</h2><pre>" + structured_html + "</pre></body></html>"
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write(body)

def main(query: str, final_text: str = "", code_result: str = "") -> dict:
    os.makedirs(EXPORT_DIR, exist_ok=True)
    fmt = _pick_format(query)
    title = _safe_name(FILENAME_HINT)
    content = _clean_text(final_text)
    suffix = "html" if fmt == "pdf" else fmt
    if suffix not in ("docx", "xlsx", "csv", "html"):
        suffix = "docx"
    filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.{suffix}"
    path = os.path.join(EXPORT_DIR, filename)
    status = "success"
    warning = ""
    try:
        if suffix == "docx":
            _write_docx(path, title, content, code_result)
        elif suffix == "xlsx":
            actual = _write_xlsx(path, title, content, code_result)
            path = actual
            filename = os.path.basename(path)
        elif suffix == "csv":
            _write_csv(path, title, content, code_result)
        else:
            _write_html(path, title, content, code_result)
            if fmt == "pdf":
                warning = "当前代码节点未检测到稳定 PDF 生成库，已生成 HTML 文件；如平台预装 reportlab/weasyprint，可在本节点改成 PDF 输出。"
    except Exception as exc:
        status = "fallback"
        warning = "目标格式生成失败，已降级生成 TXT 文件：" + str(exc)[:200]
        filename = re.sub(r"\\.[^.]+$", ".txt", filename)
        path = os.path.join(EXPORT_DIR, filename)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
            if code_result:
                f.write("\\n\\n结构化处理结果：\\n" + str(code_result))
    file_url = path
    markdown_link = f"[点击下载 {filename}]({file_url})"
    message = "文件已生成：" + file_url
    if warning:
        message += "\\n" + warning
    return {
        "file_url": file_url,
        "file_path": path,
        "filename": filename,
        "markdown_link": markdown_link,
        "status": status,
        "message": message,
    }
"""
    return template.replace("__DEFAULT_FORMAT__", fmt).replace("__FILENAME_HINT__", hint)


def render_file_export(spec, node_id, x, y, final_ref, code_ref):
    fmt = str(spec.get("format") or "auto")
    hint = str(spec.get("filename_hint") or "文件导出")
    variables = [
        {"value_selector": ["sys", "query"], "variable": "query"},
        {"value_selector": final_ref, "variable": "final_text"},
    ]
    if code_ref:
        variables.append({"value_selector": code_ref, "variable": "code_result"})
    n = node(node_id, "code", "代码生成文件", x, y, {
        "code": file_generation_code(fmt, hint),
        "code_language": "python3",
        "outputs": {
            "file_url": {"children": None, "type": "string"},
            "file_path": {"children": None, "type": "string"},
            "filename": {"children": None, "type": "string"},
            "markdown_link": {"children": None, "type": "string"},
            "status": {"children": None, "type": "string"},
            "message": {"children": None, "type": "string"},
        },
        "variables": variables,
    }, height=54)
    return n, [node_id, "markdown_link"]


def render_router(spec, node_id, x, y, upstream_refs):
    """条件路由的判断代码节点：输出 route_key 字符串供 if-else 分支。"""
    lines = as_list(spec.get("python_lines"))
    code_text = "\n".join(lines) if lines else (
        "def main(query: str = '', file_text: str = '') -> dict:\n"
        "    return {'route_key': 'doc' if str(file_text or '').strip() else 'general'}")
    code_text = _harden_none(code_text)  # str(x).strip() -> str(x or '').strip() 防 None 误判
    desc = "条件路由：" + str(spec.get("purpose") or "按条件输出 route_key")
    try:
        compile(code_text, "<router>", "exec")
    except SyntaxError as exc:
        desc += " | SYNTAX_ERROR: " + str(exc)[:160]
    if "def main" not in code_text:
        desc += " | SYNTAX_ERROR: 缺少 def main(...) 入口函数"
    if "route_key" not in code_text:
        desc += " | SYNTAX_ERROR: 路由代码必须返回 {'route_key': ...}"
    variables = []
    for item in spec.get("inputs") or []:
        var = safe_id(item.get("variable"), "var")
        src = str(item.get("from") or "query")
        if src == "query":
            variables.append({"value_selector": ["sys", "query"], "variable": var})
        elif src in _FROM_FIELD and upstream_refs.get(src):
            variables.append({"value_selector": [upstream_refs[src], _FROM_FIELD[src]], "variable": var})
    if not variables:
        variables = [{"value_selector": ["sys", "query"], "variable": "query"}]
    n = node(node_id, "code", "条件路由判断", x, y, {
        "desc": desc, "code": code_text, "code_language": "python3",
        "outputs": {"route_key": {"children": None, "type": "string"}},
        "variables": variables,
    }, height=54)
    return n, [node_id, "route_key"]


def render_if_else(node_id, x, y, router_id, case_values):
    """case_values: 有序的 route_key 字符串列表；每个生成一个 case，handle=case_id=该字符串，另有 'false' 兜底。"""
    cases = []
    for val in case_values:
        cases.append({
            "case_id": val,
            "conditions": [{
                "comparison_operator": "is",
                "id": str(uuid.uuid4()),
                "value": val,
                "varType": "string",
                "variable_selector": [router_id, "route_key"],
            }],
            "id": val,
            "logical_operator": "and",
        })
    return node(node_id, "if-else", "条件路由", x, y,
                {"cases": cases, "logical_operator": "or"},
                height=max(140, 90 + len(cases) * 40))


def _normalize_route_nodes(route_nodes):
    """保证一条路线以 answer 结尾且 answer 之后无多余节点(否则会出现悬空/断线边)。"""
    items = [n for n in (route_nodes or []) if isinstance(n, dict) and n.get("kind")]
    out = []
    for n in items:
        out.append(n)
        if str(n.get("kind")) == "answer":
            break  # answer 是终点,丢弃其后的所有节点
    if not out or str(out[-1].get("kind")) != "answer":
        out.append({"kind": "answer", "template": "自动"})
    return out


def _resolve_answer_text(spec, refs):
    """answer 节点的输出引用由 Python 拥有（LLM 不知道真实节点ID）。
    优先级：文件导出 > 最近 LLM > HTTP > 代码 > LLM 写的模板兜底。"""
    exp = refs.get("prev_export")
    llm = refs.get("prev_llm")
    if exp:
        text = ("已生成文件：\n{{#" + exp + ".markdown_link#}}\n\n"
                "文件地址：{{#" + exp + ".file_url#}}\n"
                "文件名：{{#" + exp + ".filename#}}\n"
                "状态：{{#" + exp + ".status#}}\n\n{{#" + exp + ".message#}}")
        if llm:
            text += "\n\n内容预览：\n{{#" + llm + ".text#}}"
        return text
    if llm:
        return "{{#" + llm + ".text#}}"
    if refs.get("prev_http"):
        return "{{#" + refs["prev_http"] + ".body#}}"
    if refs.get("prev_code"):
        return "{{#" + refs["prev_code"] + ".result#}}"
    return None


def render_pipeline(route_nodes, start_id, entry_id, entry_handle, entry_type, base_x, base_y, ctx):
    nodes = []
    edges = []
    x = base_x
    y = base_y
    cur_id = entry_id
    cur_handle = entry_handle
    cur_type = entry_type
    refs = {"prev_llm": None, "prev_http": None, "prev_doc": None, "prev_code": None,
            "prev_knowledge": None, "prev_export": None}
    counter = [0]
    # entry_handle 在多路由时是每条路线唯一的分类器 handle，用它做前缀避免跨路线 ID 冲突。
    prefix = safe_id(str(start_id) + "_" + str(entry_handle), "route")

    def new_id(seed):
        counter[0] += 1
        return f"{prefix}_n{counter[0]}_{seed}"

    items = _normalize_route_nodes(route_nodes)
    for spec in items:
        kind = str(spec.get("kind") or "").strip()
        if kind == "knowledge":
            nid_ = new_id("kb")
            n, ref = render_knowledge(spec, nid_, x, y, ctx["provider"])
            refs["prev_knowledge"] = nid_
        elif kind == "document-extractor":
            nid_ = new_id("doc")
            n, ref = render_document_extractor(nid_, x, y, start_id)
            refs["prev_doc"] = nid_
        elif kind == "http":
            nid_ = new_id("http")
            n, ref = render_http(spec, nid_, x, y)
            refs["prev_http"] = nid_
        elif kind == "code":
            nid_ = new_id("code")
            n, ref = render_code(spec, nid_, x, y, refs)
            refs["prev_code"] = nid_
        elif kind == "llm":
            nid_ = new_id("llm")
            ctx_ref = [refs["prev_knowledge"], "result"] if refs["prev_knowledge"] else None
            n, ref = render_llm(spec, nid_, x, y, ctx["provider"], ctx["model_name"], ctx["memory"], ctx_ref, refs, start_id)
            refs["prev_llm"] = nid_
        elif kind == "file-export":
            nid_ = new_id("file")
            final_ref = [refs["prev_llm"], "text"] if refs["prev_llm"] else ["sys", "query"]
            code_ref = [refs["prev_code"], "result"] if refs["prev_code"] else None
            n, ref = render_file_export(spec, nid_, x, y, final_ref, code_ref)
            refs["prev_export"] = nid_
        elif kind == "answer":
            nid_ = new_id("ans")
            n = render_answer(spec, nid_, x, y, _resolve_answer_text(spec, refs))
        else:
            continue
        nodes.append(n)
        edges.append(edge(cur_id, cur_handle, nid_, cur_type, n["data"]["type"]))
        cur_id = nid_
        cur_handle = "source"
        cur_type = n["data"]["type"]
        x += 320
    return nodes, edges


def _build_start_variables(blueprint, needs_upload):
    variables = []
    for item in blueprint.get("global_inputs") or []:
        var = safe_id(item.get("variable"), "input")
        variables.append({
            "variable": var,
            "label": item.get("label") or var,
            "type": item.get("type") or "text-input",
            "required": bool(item.get("required", False)),
            "max_length": int(item.get("max_length", 256)),
            "default": item.get("default", ""),
        })
    if needs_upload:
        variables.append({
            "variable": "upload_file",
            "label": "上传文件（可选）",
            "type": "file",
            "required": False,
            "allowed_file_types": ["image", "document"],
            "allowed_file_extensions": [".PDF", ".DOC", ".DOCX", ".XLS", ".XLSX", ".CSV", ".TXT", ".MD", ".PPT", ".PPTX", ".JPG", ".JPEG", ".PNG"],
            "allowed_file_upload_methods": ["local_file", "remote_url"],
        })
    return variables


def _extract_env_refs(text):
    return sorted(set(re.findall(r"\{\{#env\.([A-Za-z0-9_]+)#\}\}", str(text or ""))))


def _build_env_vars(blueprint, dsl_text):
    declared = {}
    order = []
    for item in blueprint.get("global_env_vars") or []:
        name = safe_id(item.get("name"), "ENV").upper()
        if name not in declared:
            order.append(name)
        declared[name] = {
            "description": item.get("description") or name,
            "id": str(uuid.uuid4()),
            "name": name,
            "value": "" if item.get("value_type") == "secret" else str(item.get("default", "")),
            "value_type": "secret" if item.get("value_type") == "secret" else "string",
        }
    for ref in _extract_env_refs(dsl_text):
        name = ref.upper()
        if name not in declared:
            order.append(name)
            is_secret = "TOKEN" in name or "SECRET" in name or "KEY" in name or "PASSWORD" in name
            declared[name] = {
                "description": name,
                "id": str(uuid.uuid4()),
                "name": name,
                "value": "",
                "value_type": "secret" if is_secret else "string",
            }
    return [declared[n] for n in order]


def _build_features(blueprint, needs_upload, has_knowledge):
    app = blueprint.get("app") or {}
    return {
        "file_upload": {
            "allowed_file_extensions": [".JPG", ".JPEG", ".PNG", ".GIF", ".WEBP", ".SVG", ".PDF", ".DOC", ".DOCX", ".XLS", ".XLSX", ".CSV", ".TXT", ".MD", ".PPT", ".PPTX"],
            "allowed_file_types": ["image", "document"],
            "allowed_file_upload_methods": ["local_file", "remote_url"],
            "enabled": bool(needs_upload),
            "fileUploadConfig": {"audio_file_size_limit": 50, "batch_count_limit": 10, "file_size_limit": 20, "image_file_size_limit": 10, "video_file_size_limit": 100, "workflow_file_upload_limit": 10},
            "image": {"enabled": bool(needs_upload), "number_limits": 3, "transfer_methods": ["local_file", "remote_url"]},
            "number_limits": 5,
        },
        "opening_statement": app.get("opening_statement") or "你好，我可以帮你处理任务。",
        "retriever_resource": {"enabled": bool(has_knowledge)},
        "sensitive_word_avoidance": {"enabled": False},
        "speech_to_text": {"enabled": False},
        "suggested_questions": (app.get("suggested_questions") or [])[:6],
        "suggested_questions_after_answer": {"enabled": True},
        "text_to_speech": {"enabled": False, "language": "", "voice": ""},
    }


def assemble(blueprint):
    app_meta = blueprint.get("app") or {}
    routing = blueprint.get("routing") or {}
    routes = [r for r in (routing.get("routes") or []) if isinstance(r, dict) and r.get("nodes")]
    provider = MODEL_PROVIDER
    model_name = (blueprint.get("model") or {}).get("name") or FLASH_MODEL_NAME
    memory = bool(blueprint.get("memory", True))
    ctx = {"provider": provider, "model_name": model_name, "memory": memory}

    all_nodes = [n for r in routes for n in (r.get("nodes") or [])]
    all_kinds = [str(n.get("kind") or "") for n in all_nodes]
    has_vision = any(str(n.get("kind")) == "llm" and n.get("vision") for n in all_nodes)
    # 文档抽取或多模态(vision)都需要开启文件上传；vision 节点还要引用 start 的 upload_file。
    needs_upload = ("document-extractor" in all_kinds) or has_vision or bool(routing.get("pre_extract"))
    has_knowledge = "knowledge" in all_kinds

    start_id = "node_start"
    nodes = [node(start_id, "start", "开始", 40, 260,
                  {"variables": _build_start_variables(blueprint, needs_upload)}, height=54)]
    edges = []

    rtype = str(routing.get("type") or "classifier")
    enabled = bool(routing.get("enabled")) and len(routes) >= 2
    if enabled and rtype == "condition":
        # 条件分流：(可选文档抽取) → 路由代码(route_key) → if-else → 各分支
        cur_id, cur_handle, cur_type = start_id, "source", "start"
        xc = 360
        doc_id = None
        if bool(routing.get("pre_extract")):
            doc_id = "node_router_doc"
            dn, _ = render_document_extractor(doc_id, xc, 240)
            nodes.append(dn)
            edges.append(edge(cur_id, cur_handle, doc_id, cur_type, "document-extractor"))
            cur_id, cur_handle, cur_type = doc_id, "source", "document-extractor"
            xc += 320
        router_id = "node_router"
        rn, _ = render_router(routing.get("router") or {}, router_id, xc, 240, {"prev_doc": doc_id})
        nodes.append(rn)
        edges.append(edge(cur_id, cur_handle, router_id, cur_type, "code"))
        xc += 320
        # 区分 else 兜底路线与普通 case 路线
        case_routes, else_route = [], None
        for r in routes:
            if r.get("is_else") or str(r.get("case")).strip().lower() == "else":
                else_route = r
            else:
                case_routes.append(r)
        if else_route is None and case_routes:
            else_route = case_routes.pop()
        case_values = [str(r.get("case") or r.get("id")) for r in case_routes]
        ife_id = "node_ifelse"
        nodes.append(render_if_else(ife_id, xc, 240, router_id, case_values))
        edges.append(edge(router_id, "source", ife_id, "code", "if-else"))
        for i, r in enumerate(case_routes, 1):
            handle = str(r.get("case") or r.get("id"))
            rn2, re2 = render_pipeline(r.get("nodes"), start_id, ife_id, handle, "if-else",
                                       xc + 320, 80 + (i - 1) * 260, ctx)
            nodes.extend(rn2)
            edges.extend(re2)
        if else_route:
            rn2, re2 = render_pipeline(else_route.get("nodes"), start_id, ife_id, "false", "if-else",
                                       xc + 320, 80 + len(case_routes) * 260, ctx)
            nodes.extend(rn2)
            edges.extend(re2)
    elif enabled:
        cls_id = "node_classifier"
        route_classes = [{"id": safe_id(r.get("id"), f"route_{i}"), "name": str(r.get("name") or f"路线{i}")[:24]}
                         for i, r in enumerate(routes, 1)]
        nodes.append(node(cls_id, "question-classifier", "问题分类器", 360, 240, {
            "classes": route_classes,
            "instruction": routing.get("classifier_instruction") or "根据用户问题选择最合适的业务路线。",
            "instructions": routing.get("classifier_instruction") or "根据用户问题选择最合适的业务路线。",
            "model": model(provider, model_name, 0.2),
            "query_variable_selector": ["sys", "query"],
            "topics": [],
            "vision": {"enabled": False},
        }, height=max(180, 90 + len(route_classes) * 34)))
        edges.append(edge(start_id, "source", cls_id, "start", "question-classifier"))
        for i, r in enumerate(routes, 1):
            handle = safe_id(r.get("id"), f"route_{i}")
            rn, re_ = render_pipeline(r.get("nodes"), start_id, cls_id, handle, "question-classifier",
                                      700, 80 + (i - 1) * 260, ctx)
            nodes.extend(rn)
            edges.extend(re_)
    else:
        first = routes[0] if routes else {"nodes": [
            {"kind": "llm", "system_prompt": "你是专业助手，请直接给出可用结果。", "user_template": "{{#sys.query#}}"},
            {"kind": "answer", "template": "{{#sys.query#}}"}]}
        rn, re_ = render_pipeline(first.get("nodes"), start_id, start_id, "source", "start", 360, 260, ctx)
        nodes.extend(rn)
        edges.extend(re_)

    app = {
        "app": {
            "description": app_meta.get("description") or "根据用户需求生成的 Dify 智能体",
            "icon": "🤖",
            "icon_background": "#E0F2FE",
            "mode": "advanced-chat",
            "name": app_meta.get("name") or "自定义智能体",
            "use_icon_as_answer_icon": False,
        },
        "dependencies": [{"current_identifier": None, "type": "package", "value": {"plugin_unique_identifier": "csg/ai_model_provider:0.1.3@744e26c8e82472f943fcc112c831a0904cc34c1b076b8cc74be3c11f03faa72f"}}],
        "kind": "app",
        "version": "0.3.0",
        "workflow": {
            "conversation_variables": [],
            "environment_variables": [],
            "features": _build_features(blueprint, needs_upload, has_knowledge),
            "graph": {"edges": edges, "nodes": nodes, "viewport": {"x": 0, "y": 80, "zoom": 0.7}},
        },
    }
    dsl_preview = dump_yaml(app)
    app["workflow"]["environment_variables"] = _build_env_vars(blueprint, dsl_preview)
    return app


def _looks_like_secret_leak(text):
    patterns = [
        r"Bearer\s+[A-Za-z0-9_\-]{16,}",
        r"token['\"]?\s*[:=]\s*['\"][A-Za-z0-9_\-]{12,}",
        r"password['\"]?\s*[:=]\s*['\"][^'\"]{6,}",
        r"secret['\"]?\s*[:=]\s*['\"][A-Za-z0-9_\-]{12,}",
    ]
    return any(re.search(p, str(text or ""), re.IGNORECASE) for p in patterns)


def validate(app, blueprint, dsl_text=""):
    graph = app["workflow"]["graph"]
    nodes = graph["nodes"]
    edges = graph["edges"]
    ids = {n["id"] for n in nodes}
    node_types = [n.get("data", {}).get("type") for n in nodes]
    warnings = []
    checks = []
    manual = []

    missing = [e for e in edges if e["source"] not in ids or e["target"] not in ids]
    if missing:
        warnings.append("存在断线边，请检查节点连接。")
    else:
        checks.append("节点连线检查通过。")
    if len(ids) != len(nodes):
        warnings.append("存在重复节点 ID。")
    else:
        checks.append("节点 ID 唯一。")

    for n in nodes:
        desc = str(n.get("data", {}).get("desc") or "")
        if "SYNTAX_ERROR" in desc:
            warnings.append("代码节点《" + n.get("data", {}).get("title", "代码处理") + "》存在 Python 语法问题，导入后请人工核对：" + desc.split("SYNTAX_ERROR:", 1)[-1].strip()[:120])

    envs = {e.get("name") for e in app["workflow"].get("environment_variables", [])}
    refs = _extract_env_refs(dsl_text)
    missing_envs = [r for r in refs if r not in envs]
    if missing_envs:
        warnings.append("存在未声明的环境变量引用：" + "、".join(missing_envs))
    elif refs:
        checks.append("环境变量引用均已声明：" + "、".join(refs))

    if _looks_like_secret_leak(dsl_text):
        warnings.append("疑似出现真实密钥，请改用环境变量。")
    else:
        checks.append("未发现明显硬编码密钥。")

    if "document-extractor" in node_types:
        checks.append("文件上传链路已包含文档文本抽取节点。")
    if any("代码生成文件" in str(n.get("data", {}).get("title")) for n in nodes):
        if ".file_url#}}" in dsl_text and ".markdown_link#}}" in dsl_text:
            checks.append("可下载文件输出链路完整。")
        else:
            warnings.append("已生成代码生成文件节点，但 Answer 未引用 file_url/markdown_link。")
        warnings.append("文件输出依赖 Dify 代码执行环境：需可写入 /tmp/dify_generated_files，并预装 python-docx；xlsx 建议预装 openpyxl。")

    routing_cfg = blueprint.get("routing") or {}
    routes = routing_cfg.get("routes") or []
    if (bool(routing_cfg.get("enabled")) and len(routes) >= 2
            and str(routing_cfg.get("type") or "classifier") != "condition"):
        answer_count = sum(1 for n in nodes if n.get("data", {}).get("type") == "answer")
        route_handles = {safe_id(r.get("id"), f"route_{i}") for i, r in enumerate(routes, 1)}
        classifier_edges = {e.get("sourceHandle") for e in edges if e.get("sourceHandle") in route_handles}
        if answer_count >= len(routes) and route_handles <= classifier_edges:
            checks.append(f"多路由结构检查通过：{len(routes)} 条分支均有独立回复。")
        else:
            warnings.append("多路由结构不完整，请检查问题分类器分支与各路线回复节点。")

    if "if-else" in node_types:
        ife = [n for n in nodes if n.get("data", {}).get("type") == "if-else"][0]
        case_ids = {c.get("case_id") for c in ife["data"].get("cases", [])}
        ife_handles = {e.get("sourceHandle") for e in edges if e.get("source") == ife["id"]}
        if case_ids <= ife_handles and "false" in ife_handles:
            checks.append("条件分流(if-else)检查通过：各 case 分支与 else 兜底均已接线。")
        else:
            warnings.append("条件分流(if-else)分支不完整：请检查路由代码 route_key 取值与各 case/else 分支。")

    if "/v1/your-endpoint" in dsl_text:
        warnings.append("HTTP 路径仍是占位 /v1/your-endpoint，正式使用前要改成真实接口。")
    if "knowledge-retrieval" in node_types:
        empty_kb = any(n["data"].get("type") == "knowledge-retrieval" and not n["data"].get("dataset_ids") for n in nodes)
        if empty_kb:
            warnings.append("知识检索节点未填知识库 ID，导入后需手动选择知识库。")
            manual.append("导入后在知识检索节点选择实际知识库。")
    if "http-request" in node_types:
        manual.append("导入后填写 HTTP 环境变量，并确认接口路径、方法和鉴权方式。")
    if any(n["data"].get("type") == "llm" and (n["data"].get("vision") or {}).get("enabled") for n in nodes):
        manual.append("已启用多模态图像识别：导入后请在“图像识别”LLM 节点选择支持视觉的多模态模型。")
        checks.append("多模态图像识别链路已配置(vision.enabled=true，读取上传图片)。")

    model_name = (blueprint.get("model") or {}).get("name") or FLASH_MODEL_NAME
    manual.append("导入后确认目标 Dify 平台已配置模型：" + model_name + "。")
    if not warnings:
        warnings.append("结构校验通过；导入后建议先预览运行，再发布。")

    uniq = lambda seq: list(dict.fromkeys([s for s in seq if s]))
    return {"warnings": uniq(warnings), "checks": uniq(checks), "manual": uniq(manual)}


def _compact_blueprint(blueprint):
    keep = {"action", "clarify", "app", "model", "memory", "design_notes", "global_env_vars", "global_inputs", "routing"}
    clean = {k: v for k, v in blueprint.items() if k in keep}
    return json.dumps(clean, ensure_ascii=False, separators=(",", ":"))


def _fallback_blueprint(query):
    return {
        "action": "generate",
        "app": {"name": "自定义智能体", "description": "根据用户需求生成的 Dify 智能体",
                "opening_statement": "你好，我可以帮你处理问答和写作任务。",
                "suggested_questions": ["你能做什么？", "帮我写一段文字"]},
        "model": {"name": FLASH_MODEL_NAME}, "memory": True,
        "global_env_vars": [], "global_inputs": [],
        "routing": {"enabled": False, "routes": [{"id": "main", "name": "主流程", "trigger": "",
            "nodes": [{"kind": "llm", "system_prompt": "你是专业助手，请直接给出可用结果。",
                       "user_template": "用户问题：{{#sys.query#}}"},
                      {"kind": "answer", "template": "{{#sys.query#}}"}]}]},
    }


def _format_clarify(blueprint):
    clarify = blueprint.get("clarify") or {}
    reason = clarify.get("reason") or "当前需求还缺少关键信息。"
    questions = as_list(clarify.get("questions")) or ["这个智能体主要用途是什么？", "需要哪些能力（知识库/文件/接口/分类）？"]
    lines = ["我先不生成 YAML，因为有关键需求不明确：", "", "- " + reason, "", "请补充："]
    lines += [f"{i}. {q}" for i, q in enumerate(questions, 1)]
    return "\n".join(lines)


def _format_generation(blueprint, dsl, app, validation):
    name = (blueprint.get("app") or {}).get("name") or "自定义智能体"
    filename = safe_filename(name)
    n_nodes = len(app["workflow"]["graph"]["nodes"])
    n_edges = len(app["workflow"]["graph"]["edges"])
    summary = f"已生成《{name}》（节点 {n_nodes}、连线 {n_edges}）。"
    # 干净交付版：摘要 + 设计要点(架构师的关键决定/假设,便于纠正) + 手动配置 + 继续修改提示。
    lines = [summary]
    notes = as_list(blueprint.get("design_notes"))
    if notes:
        lines += ["", "本次设计要点："] + [f"- {n}" for n in notes[:4]]
    if validation.get("manual"):
        lines += ["", "导入后需手动配置："] + [f"- {m}" for m in validation["manual"]]
    lines += ["", "继续修改：直接说“加上知识库 / 去掉HTTP / 改名为xxx / 导入报错xxx”，我基于当前草稿改；也可上传旧 .yml 让我继续优化。"]
    return summary, "\n".join(lines)


def main(query, blueprint_json, last_spec_json="", uploaded_dsl_text="", last_dsl_yaml=""):
    blueprint = clean_json(blueprint_json)
    if not blueprint:
        blueprint = _fallback_blueprint(query)
    if str(blueprint.get("action")) == "clarify":
        final_answer = _format_clarify(blueprint)
        return {
            "dsl_yaml": "", "draft_spec_json": _compact_blueprint(blueprint),
            "filename": safe_filename((blueprint.get("app") or {}).get("name") or "自定义智能体"),
            "summary": "需要先澄清需求。", "usage": "",
            "warnings": (blueprint.get("clarify") or {}).get("reason") or "",
            "final_answer": final_answer,
        }
    app = assemble(blueprint)
    dsl = dump_yaml(app)
    validation = validate(app, blueprint, dsl)
    summary, final_answer = _format_generation(blueprint, dsl, app, validation)
    usage = "\n".join(["使用步骤：", "1. 复制 YAML 保存为 .yml 导入 Dify。",
                       "2. 检查模型节点是否用你平台已有模型。",
                       "3. 填写环境变量（尤其 secret）。",
                       "4. 知识检索节点选择实际知识库。", "5. 先预览再发布。"])
    return {
        "dsl_yaml": dsl, "draft_spec_json": _compact_blueprint(blueprint),
        "filename": safe_filename((blueprint.get("app") or {}).get("name") or "自定义智能体"),
        "summary": summary, "usage": usage,
        "warnings": "注意事项：\n" + "\n".join("- " + w for w in validation["warnings"]),
        "final_answer": final_answer,
    }
